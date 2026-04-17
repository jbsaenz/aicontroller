#!/usr/bin/env python3
"""Build a polished .docx report from the markdown submission draft."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
PROJECT_ROOT = ROOT.parent
SOURCE_MD = ROOT / "word_submission_draft.md"
OUTPUT_DOCX = ROOT / "word_submission_draft.docx"


LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")


def markdown_links_to_text(value: str) -> str:
    return LINK_RE.sub(r"\1 (\2)", value)


def set_doc_margins(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = document.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(23, 58, 119)

    subtitle = document.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = RGBColor(80, 80, 80)

    h1 = document.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(23, 58, 119)

    h2 = document.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(28, 93, 153)

    h3 = document.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(48, 48, 48)


def set_paragraph_spacing(paragraph, after: float = 6, before: float = 0) -> None:
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.line_spacing = 1.15


def add_inline_runs(paragraph, text: str) -> None:
    text = markdown_links_to_text(text)
    idx = 0
    while idx < len(text):
        if text.startswith("**", idx):
            end = text.find("**", idx + 2)
            if end != -1:
                run = paragraph.add_run(text[idx + 2:end])
                run.bold = True
                idx = end + 2
                continue
        if text.startswith("`", idx):
            end = text.find("`", idx + 1)
            if end != -1:
                run = paragraph.add_run(text[idx + 1:end])
                run.font.name = "Consolas"
                run.font.size = Pt(10.5)
                idx = end + 1
                continue
        run = paragraph.add_run(text[idx])
        idx += 1


def is_table_separator_row(row: str) -> bool:
    compact = row.replace("|", "").replace("-", "").replace(":", "").strip()
    return compact == ""


def parse_table_rows(lines: list[str], start_idx: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start_idx
    while i < len(lines):
        row = lines[i].strip()
        if not row.startswith("|"):
            break
        if not is_table_separator_row(row):
            cells = [c.strip() for c in row.strip("|").split("|")]
            rows.append(cells)
        i += 1
    return rows, i


def make_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_runs(p, text)
            set_paragraph_spacing(p, after=2)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True

    document.add_paragraph("")


def add_page_number_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = "Page "
    run = p.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def resolve_image_path(raw_path: str) -> Path:
    cleaned = raw_path.strip().strip("`").strip('"').strip("'")
    candidate = Path(cleaned)
    if candidate.is_absolute():
        return candidate
    return (PROJECT_ROOT / candidate).resolve()


def maybe_insert_figure(
    document: Document, lines: list[str], idx: int
) -> tuple[bool, int]:
    line = lines[idx].strip()
    if "Figure " not in line:
        return False, idx

    caption_text = re.sub(r"^\*\*|\*\*$", "", line).strip()
    j = idx + 1
    while j < len(lines) and not lines[j].strip():
        j += 1

    image_path: Path | None = None
    if j < len(lines) and lines[j].strip().lower().startswith("suggested image:"):
        raw = lines[j].split(":", 1)[1].strip()
        image_path = resolve_image_path(raw)
        j += 1

    if image_path and image_path.exists():
        pic = document.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_run = pic.add_run()
        pic_run.add_picture(str(image_path), width=Inches(6.2))
        set_paragraph_spacing(pic, after=4)

    cap = document.add_paragraph(caption_text)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(cap, after=8)
    for run in cap.runs:
        run.italic = True
        run.bold = True

    return True, j


def build_document(lines: list[str]) -> Document:
    doc = Document()
    set_doc_margins(doc)
    configure_styles(doc)

    i = 0
    saw_title = False
    saw_subtitle = False

    while i < len(lines):
        raw = lines[i].rstrip()
        line = raw.strip()

        if not line:
            i += 1
            continue

        if line.startswith("|"):
            rows, next_idx = parse_table_rows(lines, i)
            make_table(doc, rows)
            i = next_idx
            continue

        figure_inserted, next_idx = maybe_insert_figure(doc, lines, i)
        if figure_inserted:
            i = next_idx
            continue

        if line.startswith("# "):
            text = line[2:].strip()
            if not saw_title:
                p = doc.add_paragraph(text, style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(p, after=6)
                saw_title = True
            else:
                p = doc.add_paragraph(text, style="Heading 1")
                set_paragraph_spacing(p, before=14, after=6)
            i += 1
            continue

        if line.startswith("## "):
            text = line[3:].strip()
            if saw_title and not saw_subtitle:
                p = doc.add_paragraph(text, style="Subtitle")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(p, after=14)
                saw_subtitle = True
            else:
                p = doc.add_paragraph(text, style="Heading 2")
                set_paragraph_spacing(p, before=10, after=4)
            i += 1
            continue

        if line.startswith("### "):
            p = doc.add_paragraph(line[4:].strip(), style="Heading 3")
            set_paragraph_spacing(p, before=8, after=3)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\.\s+", "", line))
            set_paragraph_spacing(p, after=3)
            i += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:].strip())
            set_paragraph_spacing(p, after=3)
            i += 1
            continue

        if BOLD_RE.fullmatch(line) and "Insert Figure" in line:
            p = doc.add_paragraph(style="Normal")
            add_inline_runs(p, line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, after=4)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        p = doc.add_paragraph(style="Normal")
        add_inline_runs(p, line)
        set_paragraph_spacing(p, after=6)
        i += 1

    add_page_number_footer(doc)
    return doc


def main() -> None:
    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    document = build_document(lines)
    document.save(OUTPUT_DOCX)
    print(f"Formatted document written to: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
